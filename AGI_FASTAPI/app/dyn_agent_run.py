from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from typing import Optional, List
from pydantic import BaseModel
from datetime import datetime, timedelta
import os
import markdown
import requests
from sqlalchemy.orm import Session

from database import DatabaseService, get_db
from models import DynamicAgent
import logging

router = APIRouter(prefix="/dyn_agent_run", tags=["dyn_agent_run"])
logger = logging.getLogger(__name__)


# Request/Response Models
class AgentEnvironmentCreate(BaseModel):
    agent_id: int
    prompt: str
    email: Optional[str] = None
    auth_code: Optional[str] = None
    token: Optional[str] = None


class AgentEnvironmentResponse(BaseModel):
    success: bool
    message: Optional[str] = None
    content: Optional[str] = None
    error: Optional[str] = None


class GoogleDriveUploadResponse(BaseModel):
    file_id: str
    file_url: str


class CalendarEventResponse(BaseModel):
    event_link: str


class GitHubRepoResponse(BaseModel):
    repos: List[dict]


# Helper Functions
def markdown_to_html(md_text: str) -> str:
    return markdown.markdown(md_text)


def generate_system_prompt(agent_details: dict) -> str:
    return f"""
    You are {agent_details['name']}, an AI agent designed to achieve the following goal:
    {agent_details['agent_goal']}.

    Your description: {agent_details['agent_description']}.

    You have to follow the instructions:{agent_details['agent_instructions']}

    Behave in a manner consistent with this goal and description. Use external tools if required 
    (e.g., {', '.join(agent_details['tools']) if agent_details['tools'] else 'none'}). 
    Provide accurate, helpful, and concise responses to user queries.
    """


# Main API Endpoints
@router.post(
    "/create",
    response_model=DynamicAgent,
    status_code=status.HTTP_201_CREATED,
    summary="Create OpenAI environment for an agent"
)
async def create_openai_environment(
        agent_id: int,
        db: Session = Depends(get_db)
):
    try:
        logger.debug(f"Starting environment creation for agent_id: {agent_id}")

        # Retrieve the agent details from the database
        logger.debug(f"Fetching agent details from database for agent_id: {agent_id}")
        db_service = DatabaseService(db)
        agent = db_service.get_agent(agent_id)
        if not agent:
            logger.error(f"Agent not found in database for agent_id: {agent_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Agent not found"
            )

        dynamic_agent = db_service.get_dynamic_agent(agent.dynamic_agent_id)
        if not dynamic_agent:
            logger.error(f"Dynamic agent not found for agent_id: {agent_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Dynamic agent configuration not found"
            )

        logger.debug(f"Retrieved agent details: {agent.__dict__}")

        # Retrieve API key from the environment table
        logger.debug(f"Fetching environment details for env_id: {agent.env_id}")
        env_details = db_service.get_environment(agent.env_id)
        if not env_details or not env_details.api_key:
            logger.error(f"No API key found for env_id: {agent.env_id}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="API key not found in environment table"
            )
        logger.debug("Environment details retrieved successfully")

        # Dynamically construct the system prompt
        agent_details = {
            "name": dynamic_agent.agent_name,
            "agent_goal": dynamic_agent.agent_goal,
            "agent_description": dynamic_agent.agent_description,
            "agent_instructions": dynamic_agent.agent_instruction,
            "tools": agent.tools.split(',') if agent.tools else []
        }
        logger.debug(f"Constructed agent details: {agent_details}")

        system_prompt = generate_system_prompt(agent_details)
        logger.debug(f"Generated system prompt: {system_prompt[:100]}...")

        # Create the OpenAI environment
        logger.debug("Initiating OpenAI API call...")
        environment_response = await call_openai_api(
            system_prompt=system_prompt,
            user_prompt=dynamic_agent.agent_description,
            api_key=env_details.api_key
        )
        logger.debug(f"OpenAI API response: {environment_response.get('response', {}).get('id', 'No ID')}")

        if environment_response.get("success"):
            logger.info(f"Successfully created environment for agent_id: {agent_id}")
            return DynamicAgent(
                id=dynamic_agent.id,
                agent_name=dynamic_agent.agent_name,
                agent_goal=dynamic_agent.agent_goal,
                agent_description=dynamic_agent.agent_description,
                agent_instructions=dynamic_agent.agent_instruction,
                email=dynamic_agent.email
            )
        else:
            error_msg = environment_response.get("error", "Failed to create OpenAI environment")
            logger.error(f"OpenAI API call failed: {error_msg}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_msg
            )

    except HTTPException as he:
        logger.error(f"HTTPException in create_openai_environment: {he.detail}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error in create_openai_environment: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )


@router.post(
    "/run",
    response_model=AgentEnvironmentResponse,
    summary="Run agent environment with user query"
)
async def run_agent_environment(
        request: AgentEnvironmentCreate,
        db: Session = Depends(get_db)
):
    global tool
    try:
        logger.info(f"Starting agent execution for agent_id: {request.agent_id}")
        logger.debug(f"Full request payload: {request.dict()}")

        # Retrieve agent details
        logger.debug(f"Fetching agent with ID: {request.agent_id}")
        db_service = DatabaseService(db)
        agent = db_service.get_agent(request.agent_id)
        if not agent:
            logger.error(f"Agent not found for ID: {request.agent_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Agent not found"
            )
        logger.debug(f"Agent details retrieved: {agent.__dict__}")

        logger.debug(f"Fetching dynamic agent with ID: {agent.dynamic_agent_id}")
        dynamic_agent = db_service.get_dynamic_agent(agent.dynamic_agent_id)
        if not dynamic_agent:
            logger.error(f"Dynamic agent not found for ID: {agent.dynamic_agent_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Dynamic agent configuration not found"
            )
        logger.debug(f"Dynamic agent details: {dynamic_agent.__dict__}")

        # Get environment details
        logger.debug(f"Fetching environment with ID: {agent.env_id}")
        env_details = db_service.get_environment(agent.env_id)
        if not env_details or not env_details.api_key:
            logger.error(f"No API key found for environment ID: {agent.env_id}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="API key not found in environment table"
            )
        logger.debug("Environment details retrieved successfully")

        # Prepare agent details
        agent_details = {
            "name": dynamic_agent.agent_name,
            "agent_goal": dynamic_agent.agent_goal,
            "agent_description": dynamic_agent.agent_description,
            "agent_instructions": dynamic_agent.agent_instruction,
            "tools": agent.tools.split(',') if agent.tools else []
        }
        logger.debug(f"Prepared agent details: {agent_details}")

        # Generate system prompt
        system_prompt = generate_system_prompt(agent_details)
        logger.debug(f"System prompt generated (length: {len(system_prompt)} chars)")

        # Call OpenAI API
        logger.info("Initiating OpenAI API call...")
        openai_response = await call_openai_api(
            system_prompt=system_prompt,
            user_prompt=request.prompt,
            api_key=env_details.api_key
        )
        logger.debug(f"OpenAI raw response: {openai_response}")

        if not openai_response.get("success"):
            error_msg = openai_response.get("error", "OpenAI API call failed")
            logger.error(f"OpenAI API call failed: {error_msg}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_msg
            )

        chat_content = openai_response.get("choices", [{}])[0].get("message", {}).get("content", "")
        logger.debug(f"Received chat content (length: {len(chat_content)} chars)")
        markdown_content = markdown_to_html(chat_content)
        logger.debug("Converted content to markdown")

        # Handle tools if specified
        if agent.tools:
            tools = agent.tools.split(',')
            logger.debug(f"Tools to process: {tools}")
            file_path = None

            try:
                # Create a temporary file for tools that need it
                if any(tool in tools for tool in ["send_to_drive", "google_calendar"]):
                    logger.debug("Creating temporary file for tool processing...")
                    file_path = await create_temp_file(chat_content)
                    logger.debug(f"Temporary file created at: {file_path}")

                # Process each tool
                for tool in tools:
                    logger.debug(f"Processing tool: {tool}")

                    if tool == "send_to_drive" and request.auth_code:
                        logger.info("Starting Google Drive upload...")
                        drive_response = await upload_to_drive(file_path, request.auth_code)
                        if not drive_response:
                            logger.error("Google Drive upload failed")
                        else:
                            logger.info(f"Google Drive upload successful. File ID: {drive_response.file_id}")

                    elif tool == "google_calendar" and request.auth_code:
                        logger.info("Creating Google Calendar event...")
                        event_details = {
                            'summary': f"Event from {agent_details['name']}",
                            'description': chat_content,
                            'start_time': datetime.utcnow() + timedelta(hours=1),
                            'end_time': datetime.utcnow() + timedelta(hours=3)
                        }
                        calendar_response = await create_google_calendar_event(event_details, request.auth_code)
                        if not calendar_response:
                            logger.error("Google Calendar event creation failed")
                        else:
                            logger.info(f"Calendar event created: {calendar_response.event_link}")

                    elif tool == "github" and request.auth_code:
                        logger.info("Fetching GitHub repositories...")
                        repos = await list_github_repositories(request.auth_code)
                        if not repos:
                            logger.error("GitHub repository fetch failed")
                        else:
                            logger.info(f"Found {len(repos.repos)} repositories")

            except Exception as tool_ex:
                logger.error(f"Error processing tool {tool}: {str(tool_ex)}", exc_info=True)

            finally:
                # Clean up temporary file
                if file_path and os.path.exists(file_path):
                    logger.debug(f"Cleaning up temporary file: {file_path}")
                    os.remove(file_path)
                    logger.debug("Temporary file removed")

        logger.info("Agent execution completed successfully")
        return {
            "success": True,
            "message": "Response processed successfully",
            "content": markdown_content
        }

    except HTTPException as he:
        logger.error(f"HTTPException in run_agent_environment: {he.detail}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error in run_agent_environment: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )


# Service Functions
async def call_openai_api(system_prompt: str, user_prompt: str, api_key: str):
    try:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": 1500
        }

        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()

        return {
            "success": True,
            "response": response.json()
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


async def create_temp_file(content: str) -> str:
    from docx import Document
    file_name = "response.docx"
    file_path = f"./{file_name}"

    document = Document()
    document.add_heading("Response Content", level=1)
    document.add_paragraph(content)
    document.save(file_path)

    return file_path


async def upload_to_drive(file_path: str, authorization_code: str) -> Optional[GoogleDriveUploadResponse]:
    try:
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaFileUpload

        CLIENT_SECRETS_FILE = "google_credentials.json"
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"

        if not os.path.exists(CLIENT_SECRETS_FILE):
            logger.error(f"Client secrets file not found at {CLIENT_SECRETS_FILE}")
            return None

        if not os.path.exists(file_path):
            logger.error(f"File to upload not found at {file_path}")
            return None

        # Exchange authorization code for tokens
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=["https://www.googleapis.com/auth/drive.file"],
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(code=authorization_code)

        # Upload file
        drive_service = build('drive', 'v3', credentials=flow.credentials)
        file_metadata = {'name': os.path.basename(file_path)}
        media = MediaFileUpload(file_path,
                                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        drive_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,webViewLink'
        ).execute()

        return GoogleDriveUploadResponse(
            file_id=drive_file.get('id'),
            file_url=drive_file.get('webViewLink')
        )

    except Exception as e:
        logger.error(f"Failed to upload to Google Drive: {str(e)}")
        return None


async def create_google_calendar_event(event_details: dict, authorization_code: str) -> Optional[CalendarEventResponse]:
    try:
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build

        CLIENT_SECRETS_FILE = "google_credentials.json"
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"
        SCOPES = ['https://www.googleapis.com/auth/calendar']

        if not os.path.exists(CLIENT_SECRETS_FILE):
            logger.error(f"Client secrets file not found at {CLIENT_SECRETS_FILE}")
            return None

        # Exchange authorization code for tokens
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=SCOPES,
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(code=authorization_code)

        # Create event
        service = build('calendar', 'v3', credentials=flow.credentials)
        start_time = event_details.get('start_time', datetime.utcnow())
        end_time = event_details.get('end_time', start_time + timedelta(hours=1))

        event = {
            'summary': event_details.get('summary', 'No Title'),
            'description': event_details.get('description', ''),
            'start': {'dateTime': start_time.isoformat() + 'Z', 'timeZone': 'UTC'},
            'end': {'dateTime': end_time.isoformat() + 'Z', 'timeZone': 'UTC'},
            'reminders': {
                'useDefault': False,
                'overrides': [
                    {'method': 'email', 'minutes': 24 * 60},
                    {'method': 'popup', 'minutes': 10},
                ],
            },
        }

        created_event = service.events().insert(calendarId='primary', body=event).execute()
        return CalendarEventResponse(event_link=created_event['htmlLink'])

    except Exception as e:
        logger.error(f"Failed to create calendar event: {str(e)}")
        return None


async def list_github_repositories(authorization_code: str) -> Optional[GitHubRepoResponse]:
    try:
        CLIENT_ID = os.getenv("GITHUB_CLIENT_ID")
        CLIENT_SECRET = os.getenv("GITHUB_CLIENT_SECRET")
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"

        # Exchange code for token
        token_response = requests.post(
            "https://github.com/login/oauth/access_token",
            headers={"Accept": "application/json"},
            data={
                "client_id": CLIENT_ID,
                "client_secret": CLIENT_SECRET,
                "code": authorization_code,
                "redirect_uri": REDIRECT_URI,
            },
        )
        token_response.raise_for_status()
        access_token = token_response.json().get("access_token")

        if not access_token:
            logger.error("No access token in GitHub response")
            return None

        # Get repositories
        repos_response = requests.get(
            "https://api.github.com/user/repos",
            headers={"Authorization": f"Bearer {access_token}"}
        )
        repos_response.raise_for_status()

        repos = [{"name": repo["name"], "url": repo["html_url"]} for repo in repos_response.json()]
        return GitHubRepoResponse(repos=repos)

    except Exception as e:
        logger.error(f"Failed to fetch GitHub repos: {str(e)}")
        return None